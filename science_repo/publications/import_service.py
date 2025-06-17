import os
import tempfile
import logging
from io import BytesIO
from django.conf import settings
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
import re
import subprocess
from lxml import etree

# Try to import optional dependencies
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from pdfminer.high_level import extract_text as pdf_extract_text
except ImportError:
    pdf_extract_text = None

try:
    import docx
except ImportError:
    docx = None

logger = logging.getLogger(__name__)

class ImportService:
    """
    Service for importing manuscripts in various formats (Word, LaTeX, PDF)
    and converting them to JATS-XML.
    """

    @staticmethod
    def import_document(file_obj, file_name, document_version=None):
        """
        Import a document file and convert it to JATS-XML.

        Args:
            file_obj: The file object to import
            file_name: The name of the file
            document_version: Optional DocumentVersion instance to update

        Returns:
            dict: A dictionary containing the extracted content and metadata
        """
        file_ext = os.path.splitext(file_name)[1].lower()

        # Store the original file
        if document_version:
            path = f'documents/{document_version.id}/original{file_ext}'
            saved_path = default_storage.save(path, ContentFile(file_obj.read()))
            file_obj.seek(0)  # Reset file pointer after reading

            # Update document version with the original file path
            document_version.original_file = saved_path
            document_version.save(update_fields=['original_file'])

        # Process based on file type
        if file_ext in ['.docx', '.doc']:
            return ImportService._process_word_document(file_obj)
        elif file_ext == '.pdf':
            return ImportService._process_pdf_document(file_obj)
        elif file_ext in ['.tex', '.latex']:
            return ImportService._process_latex_document(file_obj, file_name)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

    @staticmethod
    def _process_word_document(file_obj):
        """Process a Word document and extract content and metadata."""
        # Save to a temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file.write(file_obj.read())
            temp_path = temp_file.name

        try:
            # Parse the document
            doc = docx.Document(temp_path)

            # Extract content
            content = {
                'title': doc.paragraphs[0].text if doc.paragraphs else '',
                'abstract': '',
                'sections': [],
                'references': [],
                'authors': [],
                'doi': '',
                'orcid_ids': [],
                'funding_info': ''
            }

            # Extract abstract (usually after title and before first heading)
            abstract_text = []
            in_abstract = False

            # Process paragraphs
            current_section = {'title': '', 'content': []}

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Check if it's a heading
                if para.style.name.startswith('Heading'):
                    # Save previous section if it exists
                    if current_section['content']:
                        content['sections'].append(current_section)
                        current_section = {'title': text, 'content': []}
                    else:
                        current_section['title'] = text

                # Check for abstract
                elif not in_abstract and not content['abstract'] and 'abstract' in para.text.lower():
                    in_abstract = True
                    abstract_text.append(text.replace('Abstract', '').strip())

                # Check for references
                elif 'references' in para.text.lower() or 'bibliography' in para.text.lower():
                    # Start collecting references
                    in_references = True
                else:
                    if in_abstract:
                        abstract_text.append(text)
                    else:
                        current_section['content'].append(text)

            # Add the last section
            if current_section['content']:
                content['sections'].append(current_section)

            # Set abstract
            content['abstract'] = ' '.join(abstract_text)

            # Extract authors, DOI, ORCID IDs, and funding info
            content.update(ImportService._extract_metadata_from_text('\n'.join([p.text for p in doc.paragraphs])))

            # Convert to JATS-XML
            jats_xml = ImportService._convert_to_jats_xml(content)
            content['jats_xml'] = jats_xml

            return content

        finally:
            # Clean up
            os.unlink(temp_path)

    @staticmethod
    def _process_pdf_document(file_obj):
        """Process a PDF document and extract content and metadata."""
        # Extract text using PyMuPDF
        pdf_bytes = BytesIO(file_obj.read())

        try:
            # Extract text using PyMuPDF
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text()

            # Extract content
            content = {
                'title': '',
                'abstract': '',
                'sections': [],
                'references': [],
                'authors': [],
                'doi': '',
                'orcid_ids': [],
                'funding_info': ''
            }

            # Try to extract title (usually the first line)
            lines = text.split('\n')
            if lines:
                content['title'] = lines[0].strip()

            # Try to extract abstract
            abstract_match = re.search(r'(?i)abstract[:\s]*(.*?)(?:\n\n|\n[A-Z])', text, re.DOTALL)
            if abstract_match:
                content['abstract'] = abstract_match.group(1).strip()

            # Extract sections (this is a simplified approach)
            # In a real implementation, you would need more sophisticated parsing
            section_pattern = re.compile(r'(?m)^(\d+\.?\s+[A-Z][^.\n]+)$')
            sections = section_pattern.findall(text)

            for i, section_title in enumerate(sections):
                start_pos = text.find(section_title)
                end_pos = text.find(sections[i+1]) if i < len(sections) - 1 else len(text)
                section_content = text[start_pos + len(section_title):end_pos].strip()
                content['sections'].append({
                    'title': section_title.strip(),
                    'content': [section_content]
                })

            # Extract metadata
            content.update(ImportService._extract_metadata_from_text(text))

            # Convert to JATS-XML
            jats_xml = ImportService._convert_to_jats_xml(content)
            content['jats_xml'] = jats_xml

            return content

        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            raise

    @staticmethod
    def _process_latex_document(file_obj, file_name):
        """Process a LaTeX document and extract content and metadata."""
        # Save to a temporary file
        with tempfile.NamedTemporaryFile(suffix='.tex', delete=False) as temp_file:
            temp_file.write(file_obj.read())
            temp_path = temp_file.name

        try:
            # Read the LaTeX file
            with open(temp_path, 'r', encoding='utf-8') as f:
                latex_content = f.read()

            # Extract content
            content = {
                'title': '',
                'abstract': '',
                'sections': [],
                'references': [],
                'authors': [],
                'doi': '',
                'orcid_ids': [],
                'funding_info': ''
            }

            # Extract title
            title_match = re.search(r'\\title{(.*?)}', latex_content, re.DOTALL)
            if title_match:
                content['title'] = title_match.group(1).strip()

            # Extract abstract
            abstract_match = re.search(r'\\begin{abstract}(.*?)\\end{abstract}', latex_content, re.DOTALL)
            if abstract_match:
                content['abstract'] = abstract_match.group(1).strip()

            # Extract authors
            author_match = re.search(r'\\author{(.*?)}', latex_content, re.DOTALL)
            if author_match:
                authors_text = author_match.group(1)
                # Split by \and if multiple authors
                authors = re.split(r'\\and', authors_text)
                content['authors'] = [author.strip() for author in authors]

            # Extract sections
            section_pattern = re.compile(r'\\section{(.*?)}(.*?)(?=\\section{|\\end{document}|$)', re.DOTALL)
            sections = section_pattern.findall(latex_content)

            for section_title, section_content in sections:
                content['sections'].append({
                    'title': section_title.strip(),
                    'content': [section_content.strip()]
                })

            # Extract references
            # This is a simplified approach; in reality, you'd need to parse the .bib file
            bib_pattern = re.compile(r'\\bibliography{(.*?)}')
            bib_match = bib_pattern.search(latex_content)
            if bib_match:
                bib_file = bib_match.group(1)
                # In a real implementation, you would parse the .bib file here

            # Extract DOI, ORCID IDs, and funding info
            doi_match = re.search(r'\\doi{(.*?)}', latex_content)
            if doi_match:
                content['doi'] = doi_match.group(1).strip()

            # Look for ORCID IDs in the document
            orcid_pattern = re.compile(r'\\orcid{(.*?)}')
            content['orcid_ids'] = orcid_pattern.findall(latex_content)

            # Look for funding information
            funding_match = re.search(r'\\funding{(.*?)}', latex_content)
            if funding_match:
                content['funding_info'] = funding_match.group(1).strip()

            # Convert to JATS-XML
            jats_xml = ImportService._convert_to_jats_xml(content)
            content['jats_xml'] = jats_xml

            return content

        finally:
            # Clean up
            os.unlink(temp_path)

    @staticmethod
    def _extract_metadata_from_text(text):
        """Extract metadata from text content."""
        metadata = {
            'authors': [],
            'doi': '',
            'orcid_ids': [],
            'funding_info': ''
        }

        # Extract DOI
        doi_pattern = re.compile(r'(?:DOI|doi):\s*(10\.\d{4,}(?:\.\d+)*\/\S+)')
        doi_match = doi_pattern.search(text)
        if doi_match:
            metadata['doi'] = doi_match.group(1)

        # Extract ORCID IDs
        orcid_pattern = re.compile(r'(?:ORCID|orcid)(?:\s*ID)?(?:\s*:)?\s*(\d{4}-\d{4}-\d{4}-\d{3}[\dX])')
        metadata['orcid_ids'] = orcid_pattern.findall(text)

        # Extract funding information
        funding_pattern = re.compile(r'(?:funding|grant|supported by)(?:\s*:)?\s*([^.]+)', re.IGNORECASE)
        funding_match = funding_pattern.search(text)
        if funding_match:
            metadata['funding_info'] = funding_match.group(1).strip()

        # Extract authors (simplified approach)
        # In a real implementation, you would need more sophisticated author extraction
        author_section = ""
        if "author" in text.lower():
            author_section_match = re.search(r'(?i)authors?(?:\s*:)?\s*(.*?)(?:\n\n|\n[A-Z])', text, re.DOTALL)
            if author_section_match:
                author_section = author_section_match.group(1)
                # Split by common separators
                authors = re.split(r',|\band\b', author_section)
                metadata['authors'] = [author.strip() for author in authors if author.strip()]

        return metadata

    @staticmethod
    def _convert_to_jats_xml(content):
        """Convert extracted content to JATS-XML format."""
        # Create the root element
        root = etree.Element("article")
        root.set("xmlns:xlink", "http://www.w3.org/1999/xlink")
        root.set("xmlns:mml", "http://www.w3.org/1998/Math/MathML")
        root.set("xmlns:xsi", "http://www.w3.org/2001/XMLSchema-instance")
        root.set("article-type", "research-article")

        # Create front element
        front = etree.SubElement(root, "front")

        # Journal metadata
        journal_meta = etree.SubElement(front, "journal-meta")
        journal_id = etree.SubElement(journal_meta, "journal-id")
        journal_id.set("journal-id-type", "publisher-id")
        journal_id.text = "LSD"

        # Article metadata
        article_meta = etree.SubElement(front, "article-meta")

        # Title
        title_group = etree.SubElement(article_meta, "title-group")
        article_title = etree.SubElement(title_group, "article-title")
        article_title.text = content.get('title', '')

        # Authors
        contrib_group = etree.SubElement(article_meta, "contrib-group")
        for i, author_name in enumerate(content.get('authors', [])):
            contrib = etree.SubElement(contrib_group, "contrib")
            contrib.set("contrib-type", "author")

            # Add name
            name = etree.SubElement(contrib, "name")
            # Simple name parsing (assuming "First Last" format)
            name_parts = author_name.split()
            if len(name_parts) > 1:
                surname = etree.SubElement(name, "surname")
                surname.text = name_parts[-1]
                given_names = etree.SubElement(name, "given-names")
                given_names.text = " ".join(name_parts[:-1])
            else:
                surname = etree.SubElement(name, "surname")
                surname.text = author_name

            # Add ORCID if available
            if i < len(content.get('orcid_ids', [])):
                contrib_id = etree.SubElement(contrib, "contrib-id")
                contrib_id.set("contrib-id-type", "orcid")
                contrib_id.text = content['orcid_ids'][i]

        # Abstract
        if content.get('abstract'):
            abstract = etree.SubElement(article_meta, "abstract")
            p = etree.SubElement(abstract, "p")
            p.text = content['abstract']

        # DOI
        if content.get('doi'):
            article_id = etree.SubElement(article_meta, "article-id")
            article_id.set("pub-id-type", "doi")
            article_id.text = content['doi']

        # Funding information
        if content.get('funding_info'):
            funding_group = etree.SubElement(article_meta, "funding-group")
            funding_statement = etree.SubElement(funding_group, "funding-statement")
            funding_statement.text = content['funding_info']

        # Create body element
        body = etree.SubElement(root, "body")

        # Add sections
        for section in content.get('sections', []):
            sec = etree.SubElement(body, "sec")
            title = etree.SubElement(sec, "title")
            title.text = section.get('title', '')

            for paragraph in section.get('content', []):
                p = etree.SubElement(sec, "p")
                p.text = paragraph

        # Create back element
        back = etree.SubElement(root, "back")

        # Add references
        if content.get('references'):
            ref_list = etree.SubElement(back, "ref-list")
            for i, reference in enumerate(content['references']):
                ref = etree.SubElement(ref_list, "ref")
                ref.set("id", f"ref{i+1}")

                # Simple reference format
                mixed_citation = etree.SubElement(ref, "mixed-citation")
                mixed_citation.text = reference

        # Convert to string
        return etree.tostring(root, pretty_print=True, encoding='unicode')
